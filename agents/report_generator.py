"""Report generator — produces structured Word document using python-docx."""

import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from config import OUTPUT_DIR
from logging_config import logger

# ── Color Scheme ────────────────────────────────────────────────────
NAVY = RGBColor(27, 58, 107)
GREEN = RGBColor(10, 124, 46)
RED = RGBColor(192, 0, 0)
GREY = RGBColor(136, 136, 136)
ORANGE = RGBColor(204, 102, 0)
WHITE = RGBColor(255, 255, 255)
LIGHT_NAVY = RGBColor(220, 230, 241)
STEEL = RGBColor(138, 148, 166)


def _set_cell_shading(cell, color_hex: str):
    """Set background shading for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_horizontal_rule(doc):
    """Add a navy horizontal rule between sections."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="1B3A6B"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)


def _add_heading(doc, text: str, level: int = 1, color=NAVY, alignment=WD_ALIGN_PARAGRAPH.LEFT, size: int = 14):
    """Add a styled heading."""
    heading = doc.add_heading(level=level)
    heading.alignment = alignment
    run = heading.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return heading


def _add_paragraph(doc, text: str, bold: bool = False, italic: bool = False,
                   color=None, size: int = 11, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def _add_bordered_paragraph(doc, text: str, italic: bool = False):
    """Add a paragraph with a thin border around it."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(11)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="2" w:color="1B3A6B"/>'
        f'<w:left w:val="single" w:sz="6" w:space="2" w:color="1B3A6B"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="2" w:color="1B3A6B"/>'
        f'<w:right w:val="single" w:sz="6" w:space="2" w:color="1B3A6B"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def _add_fallback_warning(doc):
    """Add a bordered fallback warning paragraph."""
    p = doc.add_paragraph()
    run = p.add_run("⚠ Generated via fallback logic — please verify this section manually")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = ORANGE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="2" w:color="CC6600"/>'
        f'<w:left w:val="single" w:sz="6" w:space="2" w:color="CC6600"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="2" w:color="CC6600"/>'
        f'<w:right w:val="single" w:sz="6" w:space="2" w:color="CC6600"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)
    return p


def _format_option_with_statute(option_text: str, applicable_law: str) -> str:
    """Ensure finding options include statute anchor text."""
    if not option_text:
        return ""
    if option_text.strip().lower().startswith("under section"):
        return option_text
    if applicable_law:
        return f"Under {applicable_law}, {option_text}"
    return option_text


def _make_table_header_row(table, row_idx: int = 0):
    """Style a table header row with navy background."""
    row = table.rows[row_idx]
    for cell in row.cells:
        _set_cell_shading(cell, "1B3A6B")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)


def _set_table_style(table):
    """Apply consistent styling to tables."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"


def generate_dss_report(
    dispute: dict,
    extracted_facts: dict,
    arbitrability_result,
    landmark_matches: list,
    landmark_analyses: list,
    issues: list,
    legal_principles: list,
    award_framework: dict,
    adversarial_analysis: dict = None,
    generation_methods: dict = None,
) -> str:
    """Generate the full DSS Word document report. Returns filepath."""
    doc = Document()

    # ── Page Setup ──────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ── HEADER ──────────────────────────────────────────────────────
    _add_paragraph(doc, "TRADEMARK ARBITRATION", bold=True, color=NAVY,
                   size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "DECISION SUPPORT SYSTEM REPORT", bold=True, color=NAVY,
                   size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    party_a = dispute.get("party_a", "Party A")
    party_b = dispute.get("party_b", "Party B")
    trademark = dispute.get("trademark_name", "N/A")

    _add_paragraph(doc, f"IN THE MATTER OF ARBITRATION BETWEEN {party_a.upper()} AND {party_b.upper()}",
                   bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    report_date = datetime.now().strftime("%d %B %Y")
    _add_paragraph(doc, f"Report Date: {report_date}", size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"Trademark in Dispute: {trademark}", size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "FOR ARBITRATOR USE ONLY — CONFIDENTIAL", bold=True, italic=True,
                   color=RED, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    _add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # SECTION I — ARBITRABILITY DETERMINATION
    # ════════════════════════════════════════════════════════════════
    _add_heading(doc, "SECTION I — ARBITRABILITY DETERMINATION", level=1, size=14)

    is_arbitrable = arbitrability_result.is_arbitrable
    status_text = arbitrability_result.status
    status_color = GREEN if is_arbitrable else RED

    _add_paragraph(doc, status_text, bold=True, color=status_color, size=16,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Check for narrative classification warning
    warning_info = getattr(arbitrability_result, "narrative_warning", None)
    if warning_info and warning_info.get("has_disagreement", False):
        msg = warning_info.get("message", "")
        p = doc.add_paragraph()
        run_title = p.add_run("⚠️ WARNING: NARRATIVE CLASSIFICATION DISAGREEMENT\n")
        run_title.bold = True
        run_title.font.color.rgb = ORANGE
        run_title.font.size = Pt(11)

        run_msg = p.add_run(msg)
        run_msg.font.size = Pt(10.5)
        run_msg.italic = True

        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="12" w:space="4" w:color="CC6600"/>'
            f'<w:left w:val="single" w:sz="12" w:space="4" w:color="CC6600"/>'
            f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="CC6600"/>'
            f'<w:right w:val="single" w:sz="12" w:space="4" w:color="CC6600"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(12)

    # Booz Allen Test Table
    booz = arbitrability_result.booz_allen_test if arbitrability_result else {}
    booz_heading = f"A. {booz.get('test', 'Booz Allen Test')}"
    _add_heading(doc, booz_heading, level=2, size=12)
    booz_judicial = booz.get("judicial_authority", "")
    if booz_judicial:
        _add_paragraph(doc, f"Judicial Authority: {booz_judicial}", italic=True, size=10, color=STEEL)
    table = doc.add_table(rows=4, cols=2, style="Table Grid")
    _set_table_style(table)

    # Header row
    table.rows[0].cells[0].text = "Parameter"
    table.rows[0].cells[1].text = "Finding"
    _make_table_header_row(table, 0)

    table.rows[1].cells[0].text = "Right Type"
    table.rows[1].cells[1].text = booz.get("right_type_label", "")

    table.rows[2].cells[0].text = "Right in Rem Indicators"
    rem_indicators = booz.get("rem_indicators", []) or []
    table.rows[2].cells[1].text = "\n".join(f"• {ind}" for ind in rem_indicators) if rem_indicators else "None identified"

    table.rows[3].cells[0].text = "Right in Personam Indicators"
    personam_indicators = booz.get("personam_indicators", []) or []
    table.rows[3].cells[1].text = "\n".join(f"• {ind}" for ind in personam_indicators) if personam_indicators else "None identified"

    # Style first column bold
    for row in table.rows[1:]:
        for p in row.cells[0].paragraphs:
            for run in p.runs:
                run.bold = True

    doc.add_paragraph()  # spacing

    # Vidya Drolia Fourfold Test Table
    vd = arbitrability_result.vidya_drolia_test if arbitrability_result else {}
    vd_heading = f"B. {vd.get('test', 'Vidya Drolia Test')}"
    _add_heading(doc, vd_heading, level=2, size=12)
    vd_judicial = vd.get("judicial_authority", "")
    if vd_judicial:
        _add_paragraph(doc, f"Judicial Authority: {vd_judicial}", italic=True, size=10, color=STEEL)

    questions = vd.get("questions", []) or []

    table2 = doc.add_table(rows=len(questions) + 1, cols=3, style="Table Grid")
    _set_table_style(table2)

    table2.rows[0].cells[0].text = "Question"
    table2.rows[0].cells[1].text = "Answer"
    table2.rows[0].cells[2].text = "Result"
    _make_table_header_row(table2, 0)

    for i, q in enumerate(questions):
        row = table2.rows[i + 1]
        row.cells[0].text = q.get("question", "")
        row.cells[1].text = "YES" if q.get("answer") else "NO"
        passes = bool(q.get("passes"))
        row.cells[2].text = "✓ PASS" if passes else "✗ FAIL"
        # Color the result cell
        for p in row.cells[2].paragraphs:
            for run in p.runs:
                run.font.color.rgb = GREEN if passes else RED
                run.bold = True

    doc.add_paragraph()

    # Arbitration Clause
    has_clause = dispute.get("has_arbitration_clause", False)
    clause_text = "YES ✓" if has_clause else "NO ✗"
    clause_color = GREEN if has_clause else RED
    p = doc.add_paragraph()
    run = p.add_run("Arbitration Clause Present: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(clause_text)
    run.bold = True
    run.font.color.rgb = clause_color
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Reason and Recommendation
    _add_heading(doc, "Determination Reasoning", level=2, size=12)
    _add_paragraph(doc, arbitrability_result.reason, size=11)

    _add_heading(doc, "Recommendation", level=2, size=12)
    _add_paragraph(doc, arbitrability_result.recommendation, size=11)

    _add_heading(doc, "Primary Authority", level=2, size=12)
    _add_paragraph(doc, arbitrability_result.primary_authority, bold=True, size=11)

    _add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # SECTION II — LANDMARK CASE ANALYSIS MATRIX
    # ════════════════════════════════════════════════════════════════
    _add_heading(doc, "SECTION II — LANDMARK CASE ANALYSIS MATRIX", level=1, size=14)

    for idx, (landmark, analysis) in enumerate(zip(landmark_matches, landmark_analyses)):
        _add_heading(doc, f"Case {idx + 1}: {landmark.case_name}", level=2, size=12)

        # Case details table
        table3 = doc.add_table(rows=6, cols=2, style="Table Grid")
        _set_table_style(table3)

        table3.rows[0].cells[0].text = "Parameter"
        table3.rows[0].cells[1].text = "Details"
        _make_table_header_row(table3, 0)

        details = [
            ("Citation", landmark.citation),
            ("Court", landmark.court),
            ("Year", str(landmark.year)),
            ("Key Principle", landmark.principle),
            ("Relevance Score", f"{landmark.similarity_score:.2%}"),
        ]
        for i, (label, value) in enumerate(details):
            table3.rows[i + 1].cells[0].text = label
            table3.rows[i + 1].cells[1].text = value
            for p in table3.rows[i + 1].cells[0].paragraphs:
                for run in p.runs:
                    run.bold = True

        doc.add_paragraph()

        # Similarities
        _add_paragraph(doc, "Similarities with Current Dispute:", bold=True, size=11)
        similarities = analysis.get("similarities", [])
        for sim in similarities:
            doc.add_paragraph(sim, style="List Bullet")

        # Differences
        _add_paragraph(doc, "Material Differences:", bold=True, size=11)
        differences = analysis.get("differences", [])
        for diff in differences:
            doc.add_paragraph(diff, style="List Bullet")

        # Binding Force
        binding = analysis.get("binding_force", "Persuasive")
        p = doc.add_paragraph()
        run = p.add_run("Binding Force: ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(binding)
        run.bold = True
        run.font.color.rgb = NAVY
        run.font.size = Pt(11)

        doc.add_paragraph()

    _add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # SECTION III — ISSUES FOR DETERMINATION
    # ════════════════════════════════════════════════════════════════
    _add_heading(doc, "SECTION III — ISSUES FOR DETERMINATION", level=1, size=14)
    if generation_methods and generation_methods.get("issues") == "fallback":
        _add_fallback_warning(doc)

    for i, issue in enumerate(issues):
        _add_paragraph(doc, f"Issue {i + 1}: {issue}", bold=True, size=11)
        doc.add_paragraph()

    _add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # SECTION IV — APPLICABLE STATUTORY PROVISIONS AND JUDICIAL INTERPRETATIONS
    # ════════════════════════════════════════════════════════════════
    _add_heading(doc, "SECTION IV — APPLICABLE STATUTORY PROVISIONS AND JUDICIAL INTERPRETATIONS", level=1, size=14)
    if generation_methods and generation_methods.get("principles") == "fallback":
        _add_fallback_warning(doc)

    for idx, principle in enumerate(legal_principles):
        p_name = principle.get("principle_name", f"Principle {idx + 1}")
        _add_heading(doc, f"{idx + 1}. {p_name}", level=2, size=12)

        table4 = doc.add_table(rows=5, cols=2, style="Table Grid")
        _set_table_style(table4)

        table4.rows[0].cells[0].text = "Parameter"
        table4.rows[0].cells[1].text = "Details"
        _make_table_header_row(table4, 0)

        statute = principle.get("statute", "") or principle.get("authority", "")
        statute_text = principle.get("statute_text", "") or principle.get("description", "")
        judicial = principle.get("judicial_interpretation", "") or principle.get("authority", "")
        rows_data = [
            ("Statutory Provision", statute),
            ("Statutory Text", statute_text),
            ("Judicial Interpretation", judicial),
            ("Application to This Dispute", principle.get("application", "")),
        ]
        for i, (label, value) in enumerate(rows_data):
            table4.rows[i + 1].cells[0].text = label
            table4.rows[i + 1].cells[1].text = value
            for p in table4.rows[i + 1].cells[0].paragraphs:
                for run in p.runs:
                    run.bold = True

        doc.add_paragraph()

    _add_horizontal_rule(doc)

    # ════════════════════════════════════════════════════════════════
    # SECTION V — AWARD FRAMEWORK
    # ════════════════════════════════════════════════════════════════
    _add_heading(doc, "SECTION V — AWARD FRAMEWORK", level=1, size=14)
    if generation_methods and generation_methods.get("framework") == "fallback":
        _add_fallback_warning(doc)

    # A. JURISDICTION
    _add_heading(doc, "A. JURISDICTION", level=2, size=12)
    jurisdiction = award_framework.get("jurisdiction_finding", "")
    _add_paragraph(doc, jurisdiction, size=11)

    doc.add_paragraph()

    # B. FINDINGS ON EACH ISSUE
    _add_heading(doc, "B. FINDINGS ON EACH ISSUE", level=2, size=12)

    findings = award_framework.get("findings_on_issues", [])
    for finding in findings:
        issue_num = finding.get("issue_number", "")
        issue_text = finding.get("issue", "")
        _add_paragraph(doc, f"Issue {issue_num}: {issue_text}", bold=True, size=11)

        applicable_law = finding.get("applicable_law", "")
        if applicable_law:
            _add_paragraph(doc, f"Applicable Law: {applicable_law}", italic=True, size=10)

        options = finding.get("finding_options", [])
        for opt_idx, option in enumerate(options):
            option_text = option.replace("Option A: ", "", 1).replace("Option B: ", "", 1).strip()
            option_text = _format_option_with_statute(option_text, applicable_law)
            p = doc.add_paragraph()
            run = p.add_run(f"□ Option {chr(65 + opt_idx)}: ")
            run.bold = True
            run.font.size = Pt(11)
            run = p.add_run(option_text)
            run.font.size = Pt(11)

        _add_paragraph(doc, "Arbitrator's Finding: _______________________________________________", size=11)
        doc.add_paragraph()

    # C. RELIEF
    _add_heading(doc, "C. RELIEF", level=2, size=12)

    relief = award_framework.get("relief_section", {})

    # Injunction
    _add_paragraph(doc, "Injunction:", bold=True, size=11)
    p = doc.add_paragraph()
    run = p.add_run("□ GRANTED    □ REFUSED")
    run.bold = True
    run.font.size = Pt(11)
    if relief.get("injunction_guidance"):
        _add_paragraph(doc, f"Guidance: {relief['injunction_guidance']}", italic=True, size=10)
    _add_paragraph(doc, "Terms: _______________________________________________", size=11)

    doc.add_paragraph()

    # Damages
    _add_paragraph(doc, "Damages:", bold=True, size=11)
    if relief.get("damages_guidance"):
        _add_paragraph(doc, f"Guidance: {relief['damages_guidance']}", italic=True, size=10)
    _add_paragraph(doc, "Amount: ₹ _______________________________________________", size=11)

    doc.add_paragraph()

    # Costs
    _add_paragraph(doc, "Costs:", bold=True, size=11)
    if relief.get("costs_guidance"):
        _add_paragraph(doc, f"Guidance: {relief['costs_guidance']}", italic=True, size=10)
    _add_paragraph(doc, "Order: _______________________________________________", size=11)

    doc.add_paragraph()

    # D. OPERATIVE PORTION
    _add_heading(doc, "D. OPERATIVE PORTION", level=2, size=12)
    operative = award_framework.get("operative_portion_template", "")
    _add_paragraph(doc, operative, size=11)

    doc.add_paragraph()

    # SECTION VI — ADVERSARIAL LEGAL ANALYSIS (optional)
    if adversarial_analysis is not None and isinstance(adversarial_analysis, dict):
        _add_horizontal_rule(doc)
        _add_heading(doc, "SECTION VI — ADVERSARIAL LEGAL ANALYSIS", level=1, size=14)
        _add_horizontal_rule(doc)
        if generation_methods and generation_methods.get("adversarial") == "fallback":
            _add_fallback_warning(doc)

        # A. LAW IN FAVOUR OF CLAIMANT
        _add_heading(doc, "A. LAW IN FAVOUR OF CLAIMANT", level=2, size=12)
        items_for = adversarial_analysis.get("law_for_claimant", []) or []
        if not items_for:
            _add_paragraph(doc, "No provisions identified.", size=11)
        for item in items_for:
            table6a = doc.add_table(rows=3, cols=2, style="Table Grid")
            _set_table_style(table6a)
            table6a.rows[0].cells[0].text = "Statutory Provision"
            table6a.rows[0].cells[1].text = item.get("statute", "")
            table6a.rows[1].cells[0].text = "Judicial Interpretation"
            table6a.rows[1].cells[1].text = item.get("case_interpretation", "")
            table6a.rows[2].cells[0].text = "Application"
            table6a.rows[2].cells[1].text = item.get("application", "")
            for r in table6a.rows:
                for p in r.cells[0].paragraphs:
                    for run in p.runs:
                        run.bold = True
            doc.add_paragraph()

        # B. LAW AGAINST CLAIMANT
        _add_heading(doc, "B. LAW AGAINST CLAIMANT", level=2, size=12)
        items_against = adversarial_analysis.get("law_against_claimant", []) or []
        if not items_against:
            _add_paragraph(doc, "No provisions identified.", size=11)
        for item in items_against:
            table6b = doc.add_table(rows=3, cols=2, style="Table Grid")
            _set_table_style(table6b)
            table6b.rows[0].cells[0].text = "Statutory Provision"
            table6b.rows[0].cells[1].text = item.get("statute", "")
            table6b.rows[1].cells[0].text = "Judicial Interpretation"
            table6b.rows[1].cells[1].text = item.get("case_interpretation", "")
            table6b.rows[2].cells[0].text = "Application"
            table6b.rows[2].cells[1].text = item.get("application", "")
            for r in table6b.rows:
                for p in r.cells[0].paragraphs:
                    for run in p.runs:
                        run.bold = True
            doc.add_paragraph()

        # C. LEGAL OPTIONS IF LAW IS AGAINST CLAIMANT
        _add_heading(doc, "C. LEGAL OPTIONS IF LAW IS AGAINST CLAIMANT", level=2, size=12)
        items_options = adversarial_analysis.get("options_if_law_against", []) or []
        if not items_options:
            _add_paragraph(doc, "No options identified.", size=11)
        for item in items_options:
            table6c = doc.add_table(rows=4, cols=2, style="Table Grid")
            _set_table_style(table6c)
            table6c.rows[0].cells[0].text = "Legal Strategy"
            table6c.rows[0].cells[1].text = item.get("option_title", "")
            table6c.rows[1].cells[0].text = "Statutory Basis"
            table6c.rows[1].cells[1].text = item.get("statute_basis", "")
            table6c.rows[2].cells[0].text = "Case Support"
            table6c.rows[2].cells[1].text = item.get("case_support", "")
            table6c.rows[3].cells[0].text = "Strength"
            strength = item.get("strength", "")
            reasoning = item.get("reasoning", "")
            cell = table6c.rows[3].cells[1]
            cell.text = ""
            p = cell.paragraphs[0]
            run_strength = p.add_run(strength)
            run_strength.bold = True
            if strength.lower() == "strong":
                run_strength.font.color.rgb = GREEN
            elif strength.lower() == "moderate":
                run_strength.font.color.rgb = ORANGE
            elif strength.lower() == "weak":
                run_strength.font.color.rgb = RED
            if reasoning:
                p.add_run(f" — {reasoning}")
            for r in table6c.rows:
                for p in r.cells[0].paragraphs:
                    for run in p.runs:
                        run.bold = True
            doc.add_paragraph()

        # D. OVERALL LEGAL POSITION ASSESSMENT
        _add_heading(doc, "D. OVERALL LEGAL POSITION ASSESSMENT", level=2, size=12)
        overall = adversarial_analysis.get("overall_legal_position", "")
        if overall:
            _add_bordered_paragraph(doc, overall, italic=True)
        else:
            _add_bordered_paragraph(doc, "No overall assessment available.", italic=True)

        _add_horizontal_rule(doc)

    # E. SIGNATURE BLOCK
    _add_heading(doc, "E. SIGNATURE BLOCK", level=2, size=12)

    doc.add_paragraph()
    _add_paragraph(doc, "Signature: _______________________________________________", size=11)
    doc.add_paragraph()
    _add_paragraph(doc, "Name: _______________________________________________", size=11)
    doc.add_paragraph()
    _add_paragraph(doc, "Date: _______________________________________________", size=11)
    doc.add_paragraph()
    _add_paragraph(doc, "Place: _______________________________________________", size=11)

    _add_horizontal_rule(doc)

    # ── FOOTER DISCLAIMER ───────────────────────────────────────────
    doc.add_paragraph()
    _add_paragraph(
        doc,
        "This report is generated by the Trademark Arbitration Decision Support System. "
        "It is an analytical tool to assist the Arbitrator and does not constitute legal "
        "advice or a binding determination. The Arbitrator retains full and independent "
        "authority to make the final award. Powered by Indian Kanoon.",
        italic=True,
        color=GREY,
        size=8,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # ── Save Document ──────────────────────────────────────────────
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Sanitize party names for filename
    safe_a = "".join(c for c in party_a if c.isalnum() or c in " _-").strip().replace(" ", "_")[:30]
    safe_b = "".join(c for c in party_b if c.isalnum() or c in " _-").strip().replace(" ", "_")[:30]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"DSS_Report_{safe_a}_v_{safe_b}_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)

    doc.save(filepath)
    logger.info(f"Report saved: {filepath}")
    return filepath
